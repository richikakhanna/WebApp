from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas
import os

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Route to upload a file
@app.route('/')
def upload_file():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def process_file():
    if 'file' not in request.files:
        return "No file part", 400
    
    file = request.files['file']
    if file.filename == '':
        return "No selected file", 400

    if file and file.filename.endswith('.docx'):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Extract metadata
        metadata = extract_metadata(file_path)
        
        # Convert to PDF
        pdf_stream = convert_to_pdf(file_path)

        return render_template('result.html', metadata=metadata, pdf_stream=pdf_stream)
    else:
        return "Invalid file type. Please upload a .docx file.", 400

# Extract metadata from the Word file
def extract_metadata(file_path):
    doc = Document(file_path)
    for para in doc.paragraphs:
        print(para.text)  # Debugging print
    core_properties = doc.core_properties
    metadata = {
        "Title": core_properties.title,
        "Author": core_properties.author,
        "Created": core_properties.created,
        "Modified": core_properties.modified
    }
    return metadata

# Convert Word file to PDF
def convert_to_pdf(file_path):
    doc = Document(file_path)
    pdf_buffer = BytesIO()
    pdf = canvas.Canvas(pdf_buffer)
    doc = Document(file_path)
    y_position = 800  # Set the initial y position for text

    for paragraph in doc.paragraphs:
        text = paragraph.text
        pdf.drawString(50, y_position, paragraph.text)
        y_position -= 20  # Move down the page for the next line
        
        # If the content goes too far down the page, create a new page
        if y_position < 100:
            pdf.showPage()
            y_position = 800  # Reset the y position for the new page

    
    pdf.save()
    pdf_buffer.seek(0)
    return pdf_buffer

# Route to download PDF
@app.route('/download', methods=['GET'])
def download_pdf():
    pdf_stream = convert_to_pdf(request.args.get('file'))
    return send_file(pdf_stream, as_attachment=True, download_name="output.pdf", mimetype='application/pdf')

if __name__ == '__main__':
    app.run(debug=True)
    app.run(host='0.0.0.0', port=5000)
